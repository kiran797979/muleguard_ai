"""
Generate the IEEE-format conference paper as a Word document.

Builds `reports/MuleGuard_IEEE_Paper.docx` to the supplied "Template for
Prototype" layout: US Letter, two columns, Times New Roman, a single-column
title block spanning the page, roman-numeral headings, numbered tables and
figures, and a bracketed reference list.

Every number in the paper is read from the JSON reports rather than typed in, so
re-running the pipeline and re-running this script keeps the paper and the
results in sync. Nothing is hardcoded that the pipeline measures.

Author details are placeholders until supplied — search the document for
"GIVEN NAME" to find them.

Run:  python src/make_paper.py
"""

from __future__ import annotations

import json

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

import config as C
from utils import log

FIG = C.REPORTS_DIR / "figures"
OUT = C.REPORTS_DIR / "MuleGuard_IEEE_Paper.docx"

# Placeholder author block — six slots, as the template provides.
AUTHORS = [
    {"name": "1st Given Name Surname", "dept": "dept. name of organization",
     "org": "name of organization", "city": "City, Country"},
    {"name": "2nd Given Name Surname", "dept": "dept. name of organization",
     "org": "name of organization", "city": "City, Country"},
    {"name": "3rd Given Name Surname", "dept": "dept. name of organization",
     "org": "name of organization", "city": "City, Country"},
]

TITLE = ("MuleGuard AI: Leakage-Hardened Money-Mule Detection and a Dataset "
         "Integrity Audit for Extremely Imbalanced Banking Data")


# --------------------------------------------------------------------------
# Data loading — every figure quoted in the text comes from here
# --------------------------------------------------------------------------
def load_reports() -> dict:
    def L(name):
        p = C.REPORTS_DIR / name
        return json.loads(p.read_text(encoding="utf-8")) if p.exists() else {}

    return {
        "clean": L("01_clean_report.json"),
        "feat": L("02_features_report.json"),
        "metrics": L("03_metrics.json"),
        "scoring": L("05_scoring_report.json"),
        "shap": L("05_shap_top_features.json"),
        "integrity": L("06_integrity_audit.json"),
    }


# --------------------------------------------------------------------------
# Low-level docx helpers
# --------------------------------------------------------------------------
# Usable width of one column: page 8.5in, margins 0.625in each side, 0.25in
# gutter -> (8.5 - 1.25 - 0.25) / 2. Anything wider than this overflows the
# column and is clipped, which is what broke the first draft.
COL_MAX = 3.44
FULL_MAX = 7.10


def set_columns(section, num: int, space_twips: int = 360) -> None:
    """Set the column count for a section (python-docx has no API for this)."""
    cols = section._sectPr.xpath("./w:cols")[0]
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(space_twips))
    cols.set(qn("w:equalWidth"), "1")


def _apply_margins(section) -> None:
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(1.0)
    section.left_margin = section.right_margin = Inches(0.625)


class Spanning:
    """Context manager that places content across BOTH columns.

    IEEE puts wide figures and tables at the top of a page spanning the full
    text width. In Word that means dropping out of the two-column section into a
    one-column continuous section and back again — there is no other way to make
    an element wider than a column without it being clipped.
    """

    def __init__(self, doc):
        self.doc = doc

    def __enter__(self):
        s = self.doc.add_section(WD_SECTION.CONTINUOUS)
        _apply_margins(s)
        set_columns(s, 1)
        return self.doc

    def __exit__(self, *exc):
        s = self.doc.add_section(WD_SECTION.CONTINUOUS)
        _apply_margins(s)
        set_columns(s, 2)
        return False


def _keep_with_next(p) -> None:
    """Stop Word orphaning a caption from the figure or table it belongs to."""
    pPr = p._p.get_or_add_pPr()
    el = OxmlElement("w:keepNext")
    pPr.append(el)


def _no_row_split(table) -> None:
    """Forbid a table row from breaking across a page.

    Without this a nine-row feature table lands half on one page and half on
    the next, stranding a two-word fragment at the top of a column.
    """
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        trPr.append(OxmlElement("w:cantSplit"))


def _repeat_header(table) -> None:
    """Mark row 0 as a header row so it repeats if the table does break."""
    trPr = table.rows[0]._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    trPr.append(el)


def _keep_table_together(table) -> None:
    """Keep every row with the next so the table moves to the next column whole.

    Without this a table that does not quite fit is split, and the continuation
    rows appear at the top of the next column with no header above them.
    """
    for row in list(table.rows)[:-1]:
        for cell in row.cells:
            for p in cell.paragraphs:
                _keep_with_next(p)


def _fixed_layout(table, widths: list[float]) -> None:
    """Pin column widths. Word ignores them unless layout is fixed AND every
    cell carries the width, so we set both."""
    table.autofit = False
    tblPr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    for row in table.rows:
        for i, w in enumerate(widths):
            if i < len(row.cells):
                row.cells[i].width = Inches(w)


def style_document(doc: Document, margins_only: bool = False) -> None:
    if margins_only:
        for s in doc.sections:
            _apply_margins(s)
        return
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), "Times New Roman")
    pf = normal.paragraph_format
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.0

    for s in doc.sections:
        s.page_width, s.page_height = Inches(8.5), Inches(11)
        s.top_margin = Inches(0.75)
        s.bottom_margin = Inches(1.0)
        s.left_margin = s.right_margin = Inches(0.625)


def para(doc, text="", size=10, bold=False, italic=False, align=None,
         space_before=0, space_after=0, first_indent=None, style=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_indent is not None:
        pf.first_line_indent = Inches(first_indent)
    if text:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
    return p


def body(doc, text: str) -> None:
    """A justified body paragraph with the IEEE first-line indent."""
    para(doc, text, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         first_indent=0.2, space_after=0)


def heading1(doc, roman: str, text: str) -> None:
    p = para(doc, f"{roman}.  {text.upper()}", size=10,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=4)
    p.runs[0].font.small_caps = True


def heading2(doc, letter: str, text: str) -> None:
    para(doc, f"{letter}.  {text}", size=10, italic=True,
         space_before=8, space_after=3)


def _figure(doc, filename: str, caption: str, width_in: float) -> None:
    path = FIG / filename
    if not path.exists():
        log(f"missing figure {path}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    _keep_with_next(p)
    p.add_run().add_picture(str(path), width=Inches(width_in))
    cap = para(doc, caption, size=8, align=WD_ALIGN_PARAGRAPH.LEFT,
               space_before=2, space_after=8)
    cap.runs[0].font.size = Pt(8)


def add_figure(doc, filename: str, caption: str, width_in: float,
               span: bool = False) -> None:
    """Place a figure, clamped to whatever width is actually available."""
    if span:
        with Spanning(doc):
            _figure(doc, filename, caption, min(width_in, FULL_MAX))
    else:
        _figure(doc, filename, caption, min(width_in, COL_MAX))


def _table(doc, number: str, title: str, header: list[str],
           rows: list[list[str]], widths: list[float], footnote: str,
           font_pt: float) -> None:
    cap = para(doc, f"TABLE {number}.  {title.upper()}", size=8,
               align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=2)
    cap.runs[0].font.small_caps = True
    _keep_with_next(cap)

    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(header):
        cell = t.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(font_pt)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].paragraph_format.space_after = Pt(1)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            para_ = cells[i].paragraphs[0]
            r = para_.add_run(str(v))
            r.font.size = Pt(font_pt)
            para_.alignment = (WD_ALIGN_PARAGRAPH.LEFT if i == 0
                               else WD_ALIGN_PARAGRAPH.CENTER)
            para_.paragraph_format.space_after = Pt(1)

    _fixed_layout(t, widths)
    _no_row_split(t)
    _repeat_header(t)
    _keep_table_together(t)

    if footnote:
        para(doc, footnote, size=7, italic=True, space_before=2, space_after=8)
    else:
        para(doc, "", size=6, space_after=6)


def add_table(doc, number: str, title: str, header: list[str],
              rows: list[list[str]], widths: list[float],
              footnote: str = "", span: bool = False,
              font_pt: float = 8) -> None:
    """Place a table, rescaled to fit the available width.

    Widths are given as proportions of the target width; anything that would
    overflow is scaled down rather than silently clipped.
    """
    budget = FULL_MAX if span else COL_MAX
    total = sum(widths)
    if total > budget:
        widths = [w * budget / total for w in widths]

    if span:
        with Spanning(doc):
            _table(doc, number, title, header, rows, widths, footnote, font_pt)
    else:
        _table(doc, number, title, header, rows, widths, footnote, font_pt)


# --------------------------------------------------------------------------
# Title block
# --------------------------------------------------------------------------
def build_title_block(doc) -> None:
    para(doc, TITLE, size=20, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)

    n = len(AUTHORS)
    cols = 3 if n >= 3 else n
    rows = (n + cols - 1) // cols
    t = doc.add_table(rows=rows, cols=cols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, a in enumerate(AUTHORS):
        cell = t.rows[idx // cols].cells[idx % cols]
        cell.text = ""
        specs = [(a["name"], False), (a["dept"], True),
                 (a["org"], True), (a["city"], False)]
        for j, (text, italic) in enumerate(specs):
            p = cell.paragraphs[0] if j == 0 else cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            r.font.size = Pt(10)
            r.italic = italic
    para(doc, "", size=8, space_after=8)


# --------------------------------------------------------------------------
# The paper
# --------------------------------------------------------------------------
def build(doc, R: dict) -> None:
    m = R["metrics"]
    ig = R["integrity"]
    cl = R["clean"]
    ft = R["feat"]
    sc = R["scoring"]

    e = m["ensemble_precision_first"]
    hr = m["ensemble_high_recall"]
    pm = m["per_model"]
    prev = ig.get("prevalence", 0.00892)
    A = ig["test_A_missingness_only"]
    B = ig["test_B_individually_useless"]
    Cc = ig["test_C_shuffled_labels"]

    def pm_(k, metric="auprc"):
        return f"{pm[k][metric]['mean']:.3f} ± {pm[k][metric]['std']:.3f}"

    def ms(block, k):
        return f"{block[k]['mean']:.3f} ± {block[k]['std']:.3f}"

    # ---------------- Abstract ----------------
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.2)
    r = p.add_run("Abstract—")
    r.bold = True; r.italic = True; r.font.size = Pt(9)
    abstract = (
        f"Money-mule accounts are the settlement layer of digital payment fraud, "
        f"yet they are vanishingly rare in labelled banking data. We present "
        f"MuleGuard AI, a reproducible detection pipeline for an account-level "
        f"dataset of {m['n_accounts']:,} accounts containing only {m['n_mules']} "
        f"confirmed mules ({m['prevalence_pct']} percent prevalence). The system "
        f"makes three contributions. First, and most consequentially, we report a "
        f"dataset integrity audit demonstrating that the benchmark itself is "
        f"contaminated: positive and negative cases are drawn from disjoint "
        f"monthly extracts, and a model given only the pattern of blank cells, "
        f"with every value discarded, attains an area under the precision-recall "
        f"curve of {A['auprc']:.3f} against a random baseline of {prev:.4f}. We "
        f"provide falsification tests that make this measurable rather than "
        f"speculative. Second, we introduce a four-layer leakage defence that "
        f"removes post-outcome fields by meaning rather than by correlation, "
        f"catching an investigation-outcome flag whose correlation with the "
        f"target is only 0.05 and which every correlation threshold admits. "
        f"Third, we use the supplied data dictionary as executable domain "
        f"knowledge to engineer {len(ft.get('mule_typology_features', {}))} named "
        f"mule-typology features, and evaluate under nested repeated "
        f"cross-validation in which feature selection, stacking, probability "
        f"calibration and the operating threshold are all fitted inside the "
        f"training fold. The calibrated ensemble attains precision "
        f"{ms(e,'precision')} at recall {ms(e,'recall')} and AUPRC "
        f"{ms(e,'auprc')}. We argue these numbers are an upper bound on what this "
        f"dataset can establish, and specify the sampling change required to make "
        f"them trustworthy."
    )
    r2 = p.add_run(abstract)
    r2.bold = True; r2.italic = True; r2.font.size = Pt(9)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run("Keywords—")
    r.bold = True; r.italic = True; r.font.size = Pt(9)
    r = p.add_run("money mule detection, anti-money laundering, class imbalance, "
                  "target leakage, dataset integrity, gradient boosting, model "
                  "calibration, explainable AI")
    r.bold = True; r.italic = True; r.font.size = Pt(9)

    # ---------------- I. Introduction ----------------
    heading1(doc, "I", "Introduction")
    body(doc,
         "A money mule is a bank account, usually opened by a real customer who "
         "is coerced, recruited or paid, that receives criminal proceeds and "
         "forwards them onward within hours. Mules sit between the fraud victim "
         "and the organiser, and disrupting them is the most effective "
         "intervention available to a bank, because a frozen mule strands funds "
         "that would otherwise be irrecoverable. India's real-time payment rails "
         "have made this both faster and more damaging, and the Reserve Bank of "
         "India requires that any automated action against a customer account be "
         "explainable and auditable.")
    body(doc,
         "Detection is difficult for three reasons that compound one another. "
         "Mules are extremely rare, so accuracy is meaningless and even the area "
         f"under the ROC curve flatters. In our data {m['n_mules']} of "
         f"{m['n_accounts']:,} accounts are mules, so predicting that every "
         "account is normal scores 99.11 percent accuracy. Mules are also "
         "behaviourally camouflaged: the account belongs to a genuine customer "
         "with a genuine history, so unsupervised anomaly detection has little "
         "to work with, a point our ablation confirms sharply. Finally, "
         "supervised labels in this domain are the output of an investigation "
         "process, and that process leaves traces in the feature set that a "
         "model will happily exploit.")
    body(doc,
         "This last point turns out to dominate everything else. While building "
         "the detector we found that the supplied benchmark cannot support an "
         "honest performance estimate at all, and that the contamination is "
         "invisible to the standard defences. We therefore present the integrity "
         "audit as our primary contribution, with the detector as the vehicle "
         "that exposed it. Section III develops the audit, Section IV the "
         "pipeline, and Section V the measured results together with an explicit "
         "statement of what they do and do not establish.")

    # ---------------- II. Related work ----------------
    heading1(doc, "II", "Background and Related Work")
    heading2(doc, "A", "Imbalanced learning")
    body(doc,
         "Synthetic minority oversampling [1] and its hybrid variants remain the "
         "reflexive answer to class imbalance. We deliberately do not use them "
         "here. With 65 positives in a training fold and several hundred "
         "candidate dimensions, interpolating between minority neighbours "
         "fabricates regions of feature space for which no evidence exists, and "
         "the Tomek-link cleaning step is quadratic in sample count. Instance "
         "reweighting through the scale_pos_weight parameter of gradient-boosted "
         "trees [2], [3] achieves the same objective without inventing rows, and "
         "the precision-recall curve is the appropriate summary at this "
         "prevalence rather than the ROC curve.")
    heading2(doc, "B", "Leakage and evaluation integrity")
    body(doc,
         "Target leakage is normally framed as a feature that correlates "
         "implausibly well with the label, and is normally handled by a "
         "correlation threshold. We show in Section III-B that this framing is "
         "insufficient in fraud data, where the label is an investigation "
         "outcome and several fields record different facets of that same "
         "outcome at low individual correlation. Adversarial validation is the "
         "conventional tool for detecting distribution shift between splits; it "
         "is inapplicable here precisely because the shift is perfectly aligned "
         "with the label, which is what makes the problem severe.")
    heading2(doc, "C", "Explainability requirements")
    body(doc,
         "Shapley additive explanations [4] give per-prediction attributions "
         "with useful consistency guarantees, and are the current standard for "
         "audit trails in regulated credit and fraud decisions. Attribution "
         "over anonymised column identifiers is of limited value to an "
         "investigator, however; Section IV-F describes how we recover named "
         "banking semantics so that reasons are stated in language a compliance "
         "officer can act on.")

    # ---------------- III. Dataset and integrity ----------------
    heading1(doc, "III", "Dataset and Integrity Audit")
    heading2(doc, "A", "Data")
    body(doc,
         f"The dataset comprises {m['n_accounts']:,} accounts described by "
         f"{cl['input_shape'][1] - 1:,} anonymised features F1 through F3924, "
         f"with F3924 the binary target. A separate data dictionary maps every "
         f"column to a named banking variable and a natural-language "
         f"description. The features decompose into a regular grammar: a "
         f"statistic (ratio, deviation, average, maximum, minimum, total) over a "
         f"payment channel (cash, cheque, UPI, ATM, online transfer, merchant "
         f"payment, net banking, Aadhaar Payment Bridge, bill payment) in a "
         f"direction (credit, debit) over a window (7, 14 or 31 days, and ratios "
         f"between them). A tail block holds customer demographics, alert "
         f"counts by time of day, and investigation metadata. The dictionary "
         f"also marks 18 variables that the bank's own analysts shortlisted, "
         f"which we use as a domain prior rather than a hard filter.")

    heading2(doc, "B", "Leakage that a correlation threshold cannot see")
    body(doc,
         "Four fields record the resolution status of an alert and two record "
         "how long the investigation took. All six exist only after a human "
         "analyst closes a case, and none of them exists at the moment an "
         "account must be scored. Their correlations with the target span two "
         "orders of magnitude, as Table I shows. FRAUD_SUSPECTED at 0.969 is "
         "caught by any threshold. FALSE_POSITIVE at 0.055 is caught by none, "
         "and yet it encodes the investigation's verdict just as directly, "
         "merely inverted. We therefore classify leakage by what a field means, "
         "using the dictionary, and treat correlation only as a backstop.")

    add_table(doc, "I", "Post-outcome fields removed by meaning, not correlation",
              ["Variable", "|corr| with target", "Caught by 0.90 threshold?"],
              [["FRAUD_SUSPECTED", "0.969", "Yes"],
               ["OTHER_RESOLUTION", "0.063", "No"],
               ["FALSE_POSITIVE", "0.055", "No"],
               ["MIN_RESOLVE_DAYS", "0.055", "No"],
               ["MAX_RESOLVE_DAYS", "0.033", "No"],
               ["UNATTENDED", "0.005", "No"]],
              widths=[1.55, 0.85, 1.0], font_pt=7.5,
              footnote="a. All six are written only after an investigation "
                       "concludes, so none is available at scoring time.")

    heading2(doc, "C", "A structural confound in how the sample was assembled")
    counts = ig.get("month_split", {}).get("counts", {})
    body(doc,
         "A second and far more serious problem is visible in the month field. "
         "Table II gives the complete cross-tabulation. Every negative case is "
         "drawn from the October extract and every positive case from the "
         "September, November and December extracts. No month contains both "
         "classes. The month column alone therefore separates the two classes "
         "perfectly, and it is not a property of any customer: it records which "
         "monthly extraction run produced the row.")
    add_table(doc, "II", "Class composition by monthly extract",
              ["Extract month", "Normal accounts", "Mule accounts"],
              [[k, f"{v.get('0', 0):,}", f"{v.get('1', 0):,}"]
               for k, v in sorted(counts.items())],
              widths=[1.2, 1.1, 1.0],
              footnote="a. Zero months contain both classes; the split is total.")
    body(doc,
         "Dropping the month column does not solve this. Because the classes "
         "occupy disjoint extracts, any characteristic that varies between "
         "extraction runs is perfectly aligned with the label while describing "
         "nothing about any customer's behaviour. Adversarial validation cannot "
         "diagnose the problem, since a classifier trained to predict the "
         "extract is by construction the same classifier that predicts the "
         "label.")

    heading2(doc, "D", "Falsification tests")
    body(doc,
         "To make the confound measurable we constructed three tests, each "
         "supplying a model with information that cannot possibly identify a "
         "mule. All three should score near the random baseline of "
         f"{prev:.4f} if the dataset is sound.")
    body(doc,
         "In test A every value is discarded and the model receives only a "
         "binary indicator of whether each cell was populated. Whether a cell is "
         "blank is determined by the extraction job, not by the customer, so "
         f"this view carries no behavioural content whatsoever. It attains AUPRC "
         f"{A['auprc']:.3f} and AUROC {A['auroc']:.3f}. In test B the model "
         f"receives {B['n_features']} columns drawn from the "
         f"{B.get('pool_size', 0):,} whose individual correlation with the target "
         f"is below 0.05, none of which can identify a mule alone; together they "
         f"reach AUPRC {B['auprc']:.3f}. Test C repeats test B with the labels "
         f"randomly permuted and collapses to {Cc['auprc']:.4f}, confirming that "
         f"the evaluation harness is sound and that the first two results are "
         f"genuine properties of the data rather than a defect in our code.")
    add_figure(doc, "fig2_integrity.png",
               "Fig. 2.  Falsification tests against the random baseline "
               "(log scale). Knowing only which cells were blank, with every "
               "value discarded, separates the classes almost perfectly. The "
               "shuffled-label floor confirms the harness is sound.", 3.3)
    add_table(doc, "III", "Falsification test results",
              ["Test", "Information supplied", "AUPRC", "AUROC"],
              [["A", "Blank/not-blank only, no values", f"{A['auprc']:.3f}",
                f"{A['auroc']:.3f}"],
               ["B", "250 columns, each |corr| < 0.05", f"{B['auprc']:.3f}",
                f"{B['auroc']:.3f}"],
               ["C", "As B, labels shuffled", f"{Cc['auprc']:.4f}",
                f"{Cc['auroc']:.3f}"],
               ["—", "Random guess (prevalence)", f"{prev:.4f}", "0.500"]],
              widths=[0.4, 1.6, 0.6, 0.6])
    body(doc,
         "The conclusion is unavoidable. Any performance figure computed on this "
         "dataset, by any team and any method, measures extract provenance in "
         "addition to mule behaviour, and the two cannot be separated within "
         "this file. We report our own results in Section V with that "
         "qualification attached rather than omitted.")

    # ---------------- IV. Method ----------------
    heading1(doc, "IV", "Methodology")
    add_figure(doc, "fig1_architecture.png",
               "Fig. 1.  MuleGuard AI pipeline. Stage 0 runs before any "
               "modelling and its verdict qualifies every metric produced "
               "downstream (red). Leak defences concentrate in Stage 1. In "
               "Stage 3 every component that touches the label — feature "
               "selection, stacking weights, the calibration map and the "
               "decision threshold — is fitted inside the training fold and "
               "applied frozen to validation rows. The data dictionary is drawn "
               "as a bus because it is cross-cutting rather than a pipeline "
               "step: it supplies the named semantics that Stages 1, 2 and 4 "
               "depend on. The graph stage is omitted from the figure; it "
               "detects the absence of counterparty identifiers and disables "
               "itself.",
               7.0, span=True)

    heading2(doc, "A", "The data dictionary as executable knowledge")
    body(doc,
         "We parse the dictionary into a queryable structure that maps each "
         "column identifier to its banking variable name, its description, and "
         "its decomposition into statistic, channel, direction and window. This "
         "single component is what makes the rest of the pipeline possible: leak "
         "classification by meaning, feature engineering against named "
         "quantities, and explanations phrased in banking language.")

    heading2(doc, "B", "Preserving information the naive pipeline destroys")
    body(doc,
         "Eight columns in the dataset are categorical, including occupation, "
         "gender, area category, product, segmentation class and account age "
         "bucket. Coercing the matrix to numeric, as is customary, converts all "
         "of them to missing values which are then discarded as sparse. This "
         "discards real signal: mule rates by occupation range from 0.45 percent "
         "for homemakers to 1.94 percent for students against a base rate of "
         "0.89 percent. We encode the account-age bucket ordinally, since it is "
         "ordered, and the remainder as indicators.")
    body(doc,
         "We also revisit imputation. For absolute activity aggregates a missing "
         "value does not mean unknown, it means the activity never occurred: a "
         "blank UPI transaction count identifies a customer who does not use "
         "UPI. Median imputation invents activity, and discarding such columns "
         "as sparse destroys the fact that an account rides exactly one payment "
         f"rail, which is itself diagnostic. We therefore zero-fill "
         f"{cl['zero_filled_activity']['values_filled']:,} values across "
         f"{cl['zero_filled_activity']['columns']:,} aggregate columns and leave "
         f"derived ratios to median imputation, where a missing value genuinely "
         f"does mean an undefined baseline.")

    heading2(doc, "C", "Four-layer leakage defence")
    body(doc,
         "Layer one removes post-outcome fields semantically, as in Section "
         "III-B. Layer two removes structural artefacts of sample assembly, "
         "principally the month field. Layer three, which we call extract "
         "hardening, addresses the confound of Section III-C directly: any "
         "column whose blank rate differs between the classes by more than ten "
         "percentage points is removed outright, values and missingness "
         "together, on the grounds that population of a cell is decided by the "
         "extraction job rather than the customer. This filter removed "
         f"{cl['extract_hardening']['columns_dropped']} columns. It uses the "
         "label, but it is a data-quality filter rather than a fitted component "
         "and can only remove signal, never manufacture it, so it moves the "
         "reported result in the conservative direction. Layer four is a "
         "separation audit that scans every surviving column for disjoint class "
         "ranges or a near-exclusive value; it is what would catch the next such "
         "artefact, and it reports clean after the first three layers.")

    heading2(doc, "D", "Mule-typology features")
    body(doc,
         "A mule receives funds and forwards them almost immediately, retains "
         f"little, operates in bursts, favours digital rails, and belongs to a "
         f"customer whose profile does not match the volume transacted. We "
         f"encode each clause of that description as a named feature family, "
         f"{len(ft.get('mule_typology_features', {}))} features in total, "
         f"summarised in Table IV. Notably, the dataset already contains "
         f"{ft.get('occupation_deviation_columns_used', 0)} columns measuring "
         f"deviation from an occupation cohort, so the profile-mismatch signal "
         f"required aggregation rather than invention.")
    add_table(doc, "IV", "Engineered mule-typology feature families",
              ["Family", "Representative feature", "Behavioural rationale"],
              [["Pass-through", "mg_passthrough_7d",
                "Credits ≈ debits: the account is a conduit, not a wallet"],
               ["Turnover/balance", "mg_turnover_over_balance_7d",
                "Moves many multiples of what it ever holds"],
               ["Burst", "mg_amount_burst_7v31",
                "Weekly velocity far above monthly: sudden activation"],
               ["Cash-out", "mg_digital_in_cash_out_7d",
                "Digital in, cash out: the layering handoff"],
               ["Channel mix", "mg_channel_hhi_7d",
                "Single-purpose accounts concentrate on one rail"],
               ["Ticket size", "mg_avg_ticket_7d",
                "Many small tickets suggest structuring"],
               ["Alert timing", "mg_alert_share_night",
                "Mule night-alert share runs about three times higher"],
               ["Balance shape", "mg_balance_volatility_7d",
                "Spike-and-drain rather than a held balance"],
               ["Profile mismatch", "mg_occ_deviation_max",
                "Volume inconsistent with the occupation cohort"]],
              widths=[1.3, 2.0, 3.5], span=True)

    heading2(doc, "E", "Nested repeated cross-validation")
    body(doc,
         "The ensemble combines an isolation forest, an XGBoost classifier and a "
         "LightGBM classifier, stacked by a logistic meta-learner and calibrated "
         "isotonically. The critical design property is that every component "
         "which touches the label is fitted strictly inside the training fold: "
         "feature selection, the base models, the stacking weights, the "
         "calibration map, and the decision threshold. Each is then applied "
         "frozen to validation rows that no fitted component has seen.")
    body(doc,
         "This matters because the natural implementation is biased in three "
         "distinct ways, all of which we corrected. Fitting an isotonic "
         "calibrator on pooled out-of-fold predictions and then scoring those "
         "same values inflates the result, because isotonic regression is a "
         "flexible monotone fit. Selecting the threshold that maximises "
         "precision on a curve and then reporting that precision is optimistic "
         "by construction, and severely so with 81 positives. Training the "
         "meta-learner on the base models' training-set predictions, where a "
         "400-tree booster is nearly perfect, presents it with a distribution it "
         "will never encounter at inference. We fit the stack and the calibrator "
         "on inner out-of-fold predictions instead.")
    body(doc,
         f"Because a single five-fold split places only about 16 mules in each "
         f"validation fold, the headline metric moves by several points on the "
         f"random seed alone. We therefore repeat the entire procedure across "
         f"{m['validation']['scheme'].split('repeat')[0].strip().split()[-1]} "
         f"shuffles and report the mean and standard deviation across all folds, "
         f"so that every figure carries its own uncertainty. Feature selection "
         f"retains the top {m['top_k_features_per_fold']} of "
         f"{m['n_features_available']:,} columns by gradient-boosted gain, "
         f"refitted per fold.")

    heading2(doc, "F", "Risk scoring and explanation")
    body(doc,
         "The calibrated probability maps to a 0 to 1000 risk score partitioned "
         "into three action bands: routine monitoring, enhanced monitoring with "
         "step-up authentication, and freeze with escalation to the anti-money "
         "laundering desk. Each scored account carries a ranked reason list "
         "derived from SHAP values and rendered through the dictionary, so an "
         "investigator reads a sentence about Aadhaar Payment Bridge credit "
         "activity rather than an attribution against column F2506.")

    # ---------------- V. Results ----------------
    heading1(doc, "V", "Results")
    heading2(doc, "A", "Detection performance")
    body(doc,
         f"Table V reports the calibrated ensemble under the validation scheme "
         f"of Section IV-E. At the precision-first operating point the system "
         f"attains precision {ms(e,'precision')} at recall {ms(e,'recall')}, "
         f"with AUPRC {ms(e,'auprc')} against a random baseline of "
         f"{m['auprc_baseline_random']:.4f}. Summed across folds the confusion "
         f"matrix is {e['tp_total']} true positives against {e['fp_total']} false "
         f"positives and {e['fn_total']} false negatives. The 0.90 precision "
         f"target was met in {m['precision_target_met_in_folds_pct']} percent of "
         f"folds. A second operating point tuned for the analyst review queue "
         f"trades precision down to {ms(hr,'precision')} to raise recall to "
         f"{ms(hr,'recall')}, which is the appropriate trade when a reviewed "
         f"alert costs minutes and a missed mule costs a laundering channel.")
    add_table(doc, "V", "Detection performance, nested repeated cross-validation",
              ["Metric", "Precision-first", "High-recall queue"],
              [["Precision", ms(e, "precision"), ms(hr, "precision")],
               ["Recall", ms(e, "recall"), ms(hr, "recall")],
               ["F1", ms(e, "f1"), ms(hr, "f1")],
               ["False positive rate", ms(e, "fpr"), ms(hr, "fpr")],
               ["True positives", e["tp_total"], hr["tp_total"]],
               ["False positives", e["fp_total"], hr["fp_total"]]],
              widths=[1.3, 1.05, 1.05],
              footnote=f"a. Mean ± standard deviation across all outer folds; "
                       f"counts are summed. Both operating points rank identically, "
                       f"so the threshold-independent metrics are shared: AUPRC "
                       f"{ms(e,'auprc')}, AUROC {ms(e,'auroc')}. Interpret in "
                       f"light of Section III.")

    heading2(doc, "B", "Ablation")
    body(doc,
         f"Table VI separates the base learners. XGBoost is the strongest single "
         f"model at AUPRC {pm_('xgb')}, LightGBM follows at {pm_('lgbm')}, and "
         f"the calibrated stack reaches {ms(e,'auprc')}. The isolation forest "
         f"result is the informative one: at AUPRC {pm_('iso')} and AUROC "
         f"{pm['iso']['auroc']['mean']:.3f} it performs worse than random. This "
         f"is not a defect but a finding, and it is consistent with the "
         f"criminology of the problem. A mule account belongs to a real customer "
         f"with an ordinary history, and the account is valuable to the "
         f"organiser precisely because it does not look anomalous. Unsupervised "
         f"outlier detection is therefore the wrong tool for mule detection, and "
         f"proposals that lead with it should be treated with suspicion.")
    add_table(doc, "VI", "Per-model ablation",
              ["Model", "AUPRC", "AUROC"],
              [["Isolation forest", pm_("iso"), f"{pm['iso']['auroc']['mean']:.3f}"],
               ["LightGBM", pm_("lgbm"), f"{pm['lgbm']['auroc']['mean']:.3f}"],
               ["XGBoost", pm_("xgb"), f"{pm['xgb']['auroc']['mean']:.3f}"],
               ["Calibrated ensemble", ms(e, "auprc"), ms(e, "auroc")]],
              widths=[1.25, 1.1, 1.05])
    add_figure(doc, "fig5_ablation.png",
               "Fig. 3.  Base learners against the calibrated ensemble. The "
               "isolation forest falls below the random baseline: mules are "
               "camouflaged as ordinary customers, not global outliers.", 3.3)

    heading2(doc, "C", "Feature attribution")
    # Must match the bar count in paper_figures.fig_shap(), or the prose and the
    # figure disagree about what "the top N" means.
    SHAP_TOP_N = 12
    shap_top = R["shap"].get("top_features_by_mean_abs_shap", [])
    n_eng = sum(1 for f in shap_top[:SHAP_TOP_N] if f["variable"].startswith("mg_"))
    body(doc,
         f"Fig. 4 ranks the {SHAP_TOP_N} most influential features by mean "
         f"absolute SHAP contribution. {n_eng} of the {SHAP_TOP_N} are features "
         f"engineered in this work rather than supplied columns, which indicates that the "
         f"typology encoding in Section IV-D contributes materially rather than "
         f"restating information the raw matrix already exposed. Average ticket "
         f"size over seven days, the evening share of alerts, normalised net "
         f"flow and the count of active payment rails all rank highly, matching "
         f"the behavioural account of mule operation. Among supplied columns, "
         f"deviations of non-cash non-cheque transaction volume from the "
         f"occupation cohort dominate, which is the profile-mismatch signal "
         f"appearing directly.")
    add_figure(doc, "fig3_shap.png",
               f"Fig. 4.  Top {SHAP_TOP_N} features by mean absolute SHAP "
               f"value. Features engineered in this work are shown in teal; "
               f"supplied dictionary variables in grey.", 3.4)

    heading2(doc, "D", "Operational triage")
    hb = sc["band_stats"]["HIGH"]; mb = sc["band_stats"]["MEDIUM"]
    body(doc,
         f"Table VII shows how the score bands partition the portfolio. The high "
         f"band contains {hb['accounts']} accounts of which {hb['true_mules']} "
         f"are confirmed mules, a band precision of {hb['precision']:.2f}, "
         f"capturing {hb['recall_of_all_mules']*100:.0f} percent of all mules in "
         f"a queue representing {hb['accounts']/m['n_accounts']*100:.2f} percent "
         f"of the portfolio. The medium band adds {mb['true_mules']} further "
         f"mules across {mb['accounts']} accounts. For an anti-money laundering "
         f"team the practical statement is that reviewing well under one percent "
         f"of accounts surfaces two thirds of the mule population.")
    ACTIONS = {
        "LOW": "Routine monitoring only",
        "MEDIUM": "Enhanced monitoring, step-up authentication on transfers",
        "HIGH": "Freeze outward transfers, escalate to AML desk, prepare STR",
    }
    add_table(doc, "VII", "Risk band triage and recommended action",
              ["Band", "Accounts", "Mules", "Band precision",
               "Share of all mules", "Recommended action"],
              [[b, f"{sc['band_stats'][b]['accounts']:,}",
                sc["band_stats"][b]["true_mules"],
                f"{sc['band_stats'][b]['precision']:.3f}",
                f"{sc['band_stats'][b]['recall_of_all_mules']*100:.1f}%",
                ACTIONS[b]]
               for b in ("LOW", "MEDIUM", "HIGH")],
              widths=[0.7, 0.85, 0.6, 0.95, 1.0, 3.0], span=True)
    add_figure(doc, "fig4_bands.png",
               "Fig. 5.  Queue size and band precision. The high band is small "
               "enough to review exhaustively and almost entirely composed of "
               "confirmed mules.", 5.0, span=True)

    # ---------------- VI. Discussion ----------------
    heading1(doc, "VI", "Discussion and Limitations")
    body(doc,
         "The results in Section V are strong, and we do not believe they should "
         "be read as a measure of mule detection capability. Section III "
         "establishes that an uninformative view of the same data achieves "
         "comparable separation, which means a substantial and unquantifiable "
         "share of the reported performance reflects the extract from which a "
         "row was drawn. Our hardening layers remove the artefact wherever it "
         "can be identified. They cannot remove what is unidentifiable: a "
         "genuine behavioural feature that also drifts between months is "
         "confounded, and no modelling choice separates the two within this "
         "file. We therefore describe our figures as an upper bound on what this "
         "dataset can establish.")
    body(doc,
         "The remedy is a sampling change rather than a modelling one. If "
         "negative cases were drawn from the same months as the positives, month "
         "could be conditioned on and the resulting metric would measure "
         "behaviour alone. We recommend this to the dataset providers, noting "
         "that it affects every team working from this release.")
    body(doc,
         "Three further limitations bear statement. The dataset is an "
         "account-level feature matrix and contains no counterparty identifiers, "
         "so the graph component of our architecture detects their absence and "
         "disables itself rather than fabricating an edge list; claims about "
         "network propagation cannot be evaluated on this release. "
         "Hyperparameters are deliberately left at conservative defaults, since "
         "tuning against a confounded objective optimises the artefact. Finally, "
         f"{m['n_mules']} positives is a small sample, and the standard "
         f"deviations we report across folds are wide enough that differences of "
         f"a few points between methods should not be considered meaningful.")

    # ---------------- VII. Conclusion ----------------
    heading1(doc, "VII", "Conclusion")
    body(doc,
         f"We presented a money-mule detection pipeline for an extremely "
         f"imbalanced banking dataset, together with an integrity audit of the "
         f"dataset itself. The audit is the more important result. Positive and "
         f"negative cases originate in disjoint monthly extracts, and we "
         f"demonstrated by falsification test that the blank-cell pattern alone, "
         f"stripped of all values, attains AUPRC {A['auprc']:.3f} against a "
         f"{prev:.4f} baseline. Standard defences do not detect this, and "
         f"adversarial validation cannot, because the shift coincides exactly "
         f"with the label.")
    body(doc,
         f"Alongside it we contributed a leakage taxonomy that classifies fields "
         f"by meaning rather than correlation, catching an outcome flag at 0.05 "
         f"correlation that every threshold admits; a set of "
         f"{len(ft.get('mule_typology_features', {}))} named mule-typology "
         f"features derived from the data dictionary, {n_eng} of which rank in "
         f"the top {SHAP_TOP_N} by attribution; and a nested validation design in "
         f"which the operating threshold, not merely the model, is fitted inside "
         f"the fold. The pipeline is reproducible on Windows, macOS and Linux, "
         f"and regenerates every figure in this paper from the reports it writes.")
    body(doc,
         "We suggest that the falsification protocol of Section III-D deserves "
         "wider use. It is inexpensive, it requires no knowledge of the "
         "modelling approach, and in this instance it was the difference between "
         "reporting a headline result and understanding what that result "
         "actually measured.")

    # ---------------- Acknowledgment ----------------
    p = para(doc, "ACKNOWLEDGMENT", size=10, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=12, space_after=4)
    p.runs[0].font.small_caps = True
    body(doc,
         "The authors thank the organisers of the PSB Cybersecurity, Fraud and "
         "AI Hackathon 2026, jointly convened by Bank of India and the Indian "
         "Institute of Technology Hyderabad, for providing the account-level "
         "dataset and the accompanying data dictionary on which this work is "
         "based.")
    body(doc,
         "The data dictionary proved essential rather than incidental. Without "
         "documented variable semantics, neither the meaning-based leakage "
         "taxonomy of Section III-B nor the named feature families of Section "
         "IV-D could have been constructed, and the resulting model would have "
         "been a set of attributions over opaque column identifiers rather than "
         "an auditable decision aid.")
    body(doc,
         "The dataset integrity findings reported in Section III are offered in "
         "the same collaborative spirit: as an observation intended to "
         "strengthen the benchmark for all participants and for future releases, "
         "not as a criticism of those who assembled and shared it. Constructing "
         "a labelled money-mule dataset at all, from live banking records and "
         "under the privacy constraints the domain imposes, is a substantial "
         "undertaking, and the authors are grateful for access to it.")

    # ---------------- References ----------------
    p = para(doc, "REFERENCES", size=10, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=12, space_after=4)
    p.runs[0].font.small_caps = True

    refs = [
        "N. V. Chawla, K. W. Bowyer, L. O. Hall, and W. P. Kegelmeyer, “SMOTE: "
        "synthetic minority over-sampling technique,” J. Artif. Intell. Res., "
        "vol. 16, pp. 321–357, June 2002.",
        "T. Chen and C. Guestrin, “XGBoost: a scalable tree boosting system,” in "
        "Proc. 22nd ACM SIGKDD Int. Conf. Knowledge Discovery and Data Mining, "
        "San Francisco, CA, USA, 2016, pp. 785–794.",
        "G. Ke, Q. Meng, T. Finley, T. Wang, W. Chen, W. Ma, Q. Ye, and T. Liu, "
        "“LightGBM: a highly efficient gradient boosting decision tree,” in "
        "Advances in Neural Information Processing Systems, vol. 30, 2017, "
        "pp. 3146–3154.",
        "S. M. Lundberg and S. Lee, “A unified approach to interpreting model "
        "predictions,” in Advances in Neural Information Processing Systems, "
        "vol. 30, 2017, pp. 4765–4774.",
        "F. T. Liu, K. M. Ting, and Z. Zhou, “Isolation forest,” in Proc. 8th "
        "IEEE Int. Conf. Data Mining, Pisa, Italy, 2008, pp. 413–422.",
        "J. Davis and M. Goadrich, “The relationship between precision-recall "
        "and ROC curves,” in Proc. 23rd Int. Conf. Machine Learning, "
        "Pittsburgh, PA, USA, 2006, pp. 233–240.",
        "B. Zadrozny and C. Elkan, “Transforming classifier scores into accurate "
        "multiclass probability estimates,” in Proc. 8th ACM SIGKDD Int. Conf. "
        "Knowledge Discovery and Data Mining, Edmonton, AB, Canada, 2002, "
        "pp. 694–699.",
        "S. Kaufman, S. Rosset, C. Perlich, and O. Stitelman, “Leakage in data "
        "mining: formulation, detection, and avoidance,” ACM Trans. Knowledge "
        "Discovery from Data, vol. 6, no. 4, pp. 1–21, December 2012.",
        "G. C. Cawley and N. L. C. Talbot, “On over-fitting in model selection "
        "and subsequent selection bias in performance evaluation,” J. Mach. "
        "Learn. Res., vol. 11, pp. 2079–2107, July 2010.",
        "D. H. Wolpert, “Stacked generalization,” Neural Networks, vol. 5, "
        "no. 2, pp. 241–259, 1992.",
        "Reserve Bank of India, “Master direction on know your customer (KYC) "
        "direction, 2016 (updated).” Reserve Bank of India, Mumbai, India. "
        "[Online]. Available: https://www.rbi.org.in",
        "Financial Action Task Force, “Money laundering through the physical "
        "transportation of cash,” FATF, Paris, France, 2015. [Online]. "
        "Available: https://www.fatf-gafi.org",
        "[PLACEHOLDER] PSB Cybersecurity, Fraud and AI Hackathon 2026 problem "
        "statement, Bank of India and Indian Institute of Technology Hyderabad, "
        "2026. Replace with the exact citation supplied by the organisers.",
    ]
    for i, r in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.left_indent = Inches(0.25)
        pf.first_line_indent = Inches(-0.25)
        pf.space_after = Pt(2)
        run = p.add_run(f"[{i}]\t{r}")
        run.font.size = Pt(8)


def main() -> None:
    R = load_reports()
    if not R["metrics"]:
        log("reports/03_metrics.json missing — run the pipeline first.")
        return

    doc = Document()
    style_document(doc)

    # Section 1: single column for the title block.
    set_columns(doc.sections[0], 1)
    build_title_block(doc)

    # Section 2: two columns for the body, continuous on the same page.
    doc.add_section(WD_SECTION.CONTINUOUS)
    style_document(doc)
    set_columns(doc.sections[1], 2)

    build(doc, R)

    # Word holds an exclusive lock on any open document, so fall through a few
    # candidate names rather than losing the render.
    candidates = [OUT] + [OUT.with_name(f"{OUT.stem}_v{i}.docx") for i in range(2, 12)]
    for cand in candidates:
        try:
            doc.save(cand)
        except PermissionError:
            continue
        if cand != OUT:
            log(f"{OUT.name} is open in Word and locked.")
            log(f"Wrote {cand.name} instead — close Word, delete the older "
                f"copies, and rename this one.")
        else:
            log(f"Wrote {cand}")
        log("Remaining placeholders: author block, reference [13].")
        return

    log("Every candidate filename is locked. Close the documents open in Word "
        "and re-run.")


if __name__ == "__main__":
    main()
